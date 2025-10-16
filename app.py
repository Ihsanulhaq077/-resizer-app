import streamlit as st
import pandas as pd
from PIL import Image
import PyPDF2
import io
import base64
from docx import Document
import openpyxl
import os
from tempfile import NamedTemporaryFile

# Set page configuration
st.set_page_config(
    page_title="File Processing App",
    page_icon="🔄",
    layout="wide"
)

# Your Google AdMob ad unit IDs
AD_UNITS = {
    "ad1": "ca-app-pub-7084609546186009/6916984525",
    "ad2": "ca-app-pub-7084609546186009/4290821188", 
    "ad3": "ca-app-pub-7084609546186009/2438683463"
}

# Custom CSS for ads styling
st.markdown("""
<style>
    .ad-container {
        padding: 15px;
        margin: 10px 0;
        border: 2px solid #e0e0e0;
        border-radius: 8px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        text-align: center;
        color: white;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .ad-title {
        font-weight: bold;
        font-size: 1.1em;
        margin-bottom: 8px;
        color: #ffeb3b;
    }
    .ad-description {
        font-size: 0.9em;
        margin-bottom: 12px;
        opacity: 0.9;
    }
    .ad-button {
        display: block;
        padding: 10px 15px;
        margin: 8px 0;
        background-color: #4CAF50;
        color: white;
        text-decoration: none;
        border-radius: 5px;
        border: none;
        cursor: pointer;
        font-weight: bold;
        transition: all 0.3s ease;
        width: 100%;
    }
    .ad-button:hover {
        background-color: #45a049;
        transform: translateY(-2px);
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.2);
    }
    .ad-premium {
        background: linear-gradient(135deg, #ff6b6b 0%, #ee5a24 100%);
    }
    .ad-featured {
        background: linear-gradient(135deg, #00b894 0%, #00a085 100%);
    }
    .ad-special {
        background: linear-gradient(135deg, #0984e3 0%, #074b83 100%);
    }
    .ad-header {
        text-align: center;
        color: #333;
        margin: 20px 0 10px 0;
        padding: 10px;
        background: linear-gradient(90deg, #667eea, #764ba2);
        color: white;
        border-radius: 5px;
    }
    .error-message {
        background-color: #ffebee;
        color: #c62828;
        padding: 10px;
        border-radius: 5px;
        border-left: 4px solid #c62828;
        margin: 10px 0;
    }
    .success-message {
        background-color: #e8f5e8;
        color: #2e7d32;
        padding: 10px;
        border-radius: 5px;
        border-left: 4px solid #2e7d32;
        margin: 10px 0;
    }
    .pdf-preview {
        border: 2px solid #e0e0e0;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        background-color: #f8f9fa;
    }
    .file-info {
        background-color: #e3f2fd;
        padding: 10px;
        border-radius: 5px;
        margin: 5px 0;
    }
    .preview-image {
        border: 1px solid #ddd;
        border-radius: 5px;
        margin: 5px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .page-preview {
        background: white;
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 15px;
        margin: 5px;
        text-align: center;
        min-height: 100px;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    .warning-message {
        background-color: #fff3cd;
        color: #856404;
        padding: 15px;
        border-radius: 5px;
        border-left: 4px solid #ffc107;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

def get_file_type(uploaded_file):
    """Detect file type based on extension and content"""
    filename = uploaded_file.name.lower()
    
    # Image files
    image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.webp']
    if any(filename.endswith(ext) for ext in image_extensions):
        return "image"
    
    # PDF files
    if filename.endswith('.pdf'):
        return "pdf"
    
    # CSV files
    if filename.endswith('.csv'):
        return "csv"
    
    # DOCX files
    if filename.endswith('.docx'):
        return "docx"
    
    # Excel files
    excel_extensions = ['.xlsx', '.xls']
    if any(filename.endswith(ext) for ext in excel_extensions):
        return "excel"
    
    # Default to unknown
    return "unknown"

def display_left_sidebar_ads():
    """Display ads in the left sidebar"""
    st.sidebar.markdown("---")
    st.sidebar.markdown('<div class="ad-header">🚀 Premium Tools</div>', unsafe_allow_html=True)
    
    # Ad 1 - Premium Image Tools
    with st.sidebar.container():
        st.markdown('<div class="ad-container ad-premium">', unsafe_allow_html=True)
        st.markdown('<div class="ad-title">🎨 AI Image Enhancer Pro</div>', unsafe_allow_html=True)
        st.markdown('<div class="ad-description">Enhance image quality with advanced AI technology</div>', unsafe_allow_html=True)
        if st.button("✨ Try Free Trial", key="sidebar_left_ad1"):
            st.session_state.ad_clicked = AD_UNITS["ad1"]
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Ad 2 - PDF Solutions
    with st.sidebar.container():
        st.markdown('<div class="ad-container ad-featured">', unsafe_allow_html=True)
        st.markdown('<div class="ad-title">📊 PDF Business Suite</div>', unsafe_allow_html=True)
        st.markdown('<div class="ad-description">Professional PDF editing & conversion tools</div>', unsafe_allow_html=True)
        if st.button("🚀 Get Started", key="sidebar_left_ad2"):
            st.session_state.ad_clicked = AD_UNITS["ad2"]
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Ad 3 - Cloud Storage
    with st.sidebar.container():
        st.markdown('<div class="ad-container ad-special">', unsafe_allow_html=True)
        st.markdown('<div class="ad-title">☁️ Secure Cloud Storage</div>', unsafe_allow_html=True)
        st.markdown('<div class="ad-description">1TB secure storage with advanced features</div>', unsafe_allow_html=True)
        if st.button("💾 Free 30GB Offer", key="sidebar_left_ad3"):
            st.session_state.ad_clicked = AD_UNITS["ad3"]
        st.markdown('</div>', unsafe_allow_html=True)

def display_right_sidebar_ads():
    """Display ads in the right sidebar"""
    st.sidebar.markdown("---")
    st.sidebar.markdown('<div class="ad-header">💼 Business Solutions</div>', unsafe_allow_html=True)
    
    # Additional ad variations using the same ad units
    ad_variations = [
        {
            "title": "🤖 AI Document Assistant",
            "desc": "Smart document processing with AI",
            "button": "🤖 Learn More",
            "ad_unit": AD_UNITS["ad1"]
        },
        {
            "title": "📈 Data Analytics Pro",
            "desc": "Advanced data analysis platform",
            "button": "📊 View Demo", 
            "ad_unit": AD_UNITS["ad2"]
        },
        {
            "title": "🛡️ Enterprise Security",
            "desc": "Military-grade file protection",
            "button": "🔒 Secure Now",
            "ad_unit": AD_UNITS["ad3"]
        }
    ]
    
    for i, ad in enumerate(ad_variations):
        with st.sidebar.container():
            st.markdown(f'<div class="ad-container ad-{"premium" if i==0 else "featured" if i==1 else "special"}">', unsafe_allow_html=True)
            st.markdown(f'<div class="ad-title">{ad["title"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="ad-description">{ad["desc"]}</div>', unsafe_allow_html=True)
            if st.button(ad["button"], key=f"sidebar_right_ad_{i}"):
                st.session_state.ad_clicked = ad["ad_unit"]
            st.markdown('</div>', unsafe_allow_html=True)

def display_main_content_ads():
    """Display ads in the main content area"""
    col1, col2, col3 = st.columns([1, 3, 1])
    
    with col1:
        st.markdown('<div class="ad-header">🔥 Hot Deals</div>', unsafe_allow_html=True)
        
        left_ads = [
            {"text": "🎯 50% Off Premium", "ad_unit": AD_UNITS["ad1"]},
            {"text": "🚀 Free AI Tools", "ad_unit": AD_UNITS["ad2"]},
            {"text": "💎 Limited Offer", "ad_unit": AD_UNITS["ad3"]},
            {"text": "📊 Analytics Suite", "ad_unit": AD_UNITS["ad1"]}
        ]
        
        for i, ad in enumerate(left_ads):
            if st.button(ad["text"], key=f"main_left_ad_{i}", use_container_width=True):
                st.session_state.ad_clicked = ad["ad_unit"]
    
    with col3:
        st.markdown('<div class="ad-header">🛠️ Top Tools</div>', unsafe_allow_html=True)
        
        right_ads = [
            {"text": "🖼️ Image Pro Pack", "ad_unit": AD_UNITS["ad1"]},
            {"text": "📄 PDF Master", "ad_unit": AD_UNITS["ad2"]},
            {"text": "📊 Data Suite", "ad_unit": AD_UNITS["ad3"]},
            {"text": "☁️ Cloud Pro", "ad_unit": AD_UNITS["ad1"]}
        ]
        
        for i, ad in enumerate(right_ads):
            if st.button(ad["text"], key=f"main_right_ad_{i}", use_container_width=True):
                st.session_state.ad_clicked = ad["ad_unit"]
        
        # Banner ad at the bottom of right column
        st.markdown("---")
        st.markdown('<div class="ad-container ad-premium">', unsafe_allow_html=True)
        st.markdown('<div class="ad-title">⭐ File Processing Pro</div>', unsafe_allow_html=True)
        st.markdown('<div class="ad-description">Unlock all premium features</div>', unsafe_allow_html=True)
        if st.button("🎁 Special Offer", key="main_banner_ad", use_container_width=True):
            st.session_state.ad_clicked = AD_UNITS["ad2"]
        st.markdown('</div>', unsafe_allow_html=True)
    
    return col2  # Return the middle column for main content

def handle_ad_clicks():
    """Handle ad click events and display tracking"""
    if hasattr(st.session_state, 'ad_clicked') and st.session_state.ad_clicked:
        ad_unit = st.session_state.ad_clicked
        
        # Display success message
        st.sidebar.success("🎉 Thank you for your interest! Redirecting to our partner...")
        
        # Simulate ad redirect (in real implementation, this would open the actual ad)
        st.markdown(f"""
        <div style='background-color: #e8f5e8; padding: 15px; border-radius: 5px; margin: 10px 0;'>
            <h4>🚀 Advertisement Redirect</h4>
            <p><strong>Ad Unit:</strong> {ad_unit}</p>
            <p>In a production environment, this would redirect to the actual advertisement.</p>
            <p><em>Simulated click tracked for: {ad_unit}</em></p>
        </div>
        """, unsafe_allow_html=True)
        
        # Here you would typically:
        # 1. Track the click in your analytics
        # 2. Redirect to the actual ad URL
        # 3. Handle the ad monetization
        
        # For demonstration, we show the ad unit being used
        st.info(f"📊 Ad Click Tracked: {ad_unit}")
        
        # Clear the click state
        del st.session_state.ad_clicked

def validate_pdf_robust(file):
    """More robust PDF validation with multiple fallback methods"""
    original_position = file.tell()
    
    try:
        # Method 1: Try PyPDF2 with strict=False
        file.seek(0)
        try:
            pdf_reader = PyPDF2.PdfReader(file, strict=False)
            if len(pdf_reader.pages) > 0:
                return pdf_reader, "PyPDF2 (lenient mode)"
        except Exception as e:
            pass
        
        return None, "Not a valid PDF file"
            
    except Exception as e:
        return None, f"Unexpected error: {str(e)}"
    finally:
        file.seek(original_position)

def process_pdf():
    st.header("📄 PDF Processing")
    
    uploaded_file = st.file_uploader("Upload PDF", type=['pdf'], key="pdf_uploader")
    
    if uploaded_file is not None:
        # Check if uploaded file is actually a PDF
        file_type = get_file_type(uploaded_file)
        if file_type != "pdf":
            st.markdown(f"""
            <div class="warning-message">
            <h4>⚠️ Wrong File Type Detected</h4>
            <p>You uploaded a <strong>{file_type.upper()}</strong> file, but this section is for <strong>PDF</strong> files.</p>
            <p><strong>Uploaded file:</strong> {uploaded_file.name}</p>
            <p>Please switch to the <strong>{file_type.capitalize()}</strong> section or upload a PDF file.</p>
            </div>
            """, unsafe_allow_html=True)
            return
        
        try:
            # Validate PDF automatically when uploaded
            with st.spinner("🔍 Analyzing PDF file..."):
                pdf_reader, status = validate_pdf_robust(uploaded_file)
            
            if pdf_reader is None:
                st.error(f"""
                ❌ **Unable to process this PDF file.** 
                
                **Error:** {status}
                
                This might be because:
                - The PDF is corrupted or damaged
                - The PDF is password protected
                - The PDF uses advanced encryption
                - The file is not a valid PDF
                
                Please try with a different PDF file.
                """)
                return
            
            # Automatically display PDF preview and information
            display_pdf_preview(uploaded_file, pdf_reader)
            
            st.success("✅ PDF loaded successfully!")
            
            # Get page count
            page_count = len(pdf_reader.pages)
            
            st.markdown("---")
            st.subheader("🛠️ PDF Processing Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**📑 Extract Pages**")
                extract_all = st.checkbox("Extract all pages", key="extract_all", value=True)
                if not extract_all:
                    page_range = st.text_input("Page range (e.g., 1-3,5)", value="1", key="page_range")
                    st.caption("Examples: '1-3' for pages 1 to 3, '1,3,5' for specific pages")
            
            with col2:
                st.markdown("**🔄 Merge PDFs**")
                merge_files = st.file_uploader("Upload additional PDFs to merge", type=['pdf'], accept_multiple_files=True, key="merge_files")
                if merge_files:
                    st.success(f"📎 {len(merge_files)} additional PDF(s) ready for merging")
            
            # Process PDF button
            if st.button("🚀 Process PDF", key="process_pdf_btn", type="primary"):
                try:
                    with st.spinner("🔄 Processing PDF..."):
                        # Create PDF writer
                        pdf_writer = PyPDF2.PdfWriter()
                        
                        # Process main PDF
                        if extract_all:
                            for page in pdf_reader.pages:
                                pdf_writer.add_page(page)
                            st.info(f"✅ Added all {page_count} pages")
                        else:
                            pages_to_extract = parse_page_range(page_range, page_count)
                            for page_num in pages_to_extract:
                                if 1 <= page_num <= page_count:
                                    pdf_writer.add_page(pdf_reader.pages[page_num-1])
                            st.info(f"✅ Added {len(pages_to_extract)} selected pages")
                        
                        # Process merge files
                        if merge_files:
                            merge_count = 0
                            for merge_file in merge_files:
                                merge_reader, merge_status = validate_pdf_robust(merge_file)
                                if merge_reader:
                                    for page in merge_reader.pages:
                                        pdf_writer.add_page(page)
                                        merge_count += 1
                                else:
                                    st.warning(f"⚠️ Skipped one PDF: {merge_status}")
                            
                            if merge_count > 0:
                                st.info(f"✅ Merged {merge_count} additional pages")
                        
                        # Save to bytes buffer
                        output_buffer = io.BytesIO()
                        pdf_writer.write(output_buffer)
                        output_buffer.seek(0)
                        
                        # Show success message
                        st.success("🎉 PDF processed successfully!")
                        
                        # Show final document info
                        final_page_count = len(pdf_writer.pages)
                        st.info(f"📄 Final document: {final_page_count} pages")
                        
                        # Download button
                        st.download_button(
                            label="📥 Download Processed PDF",
                            data=output_buffer.getvalue(),
                            file_name="processed_pdf.pdf",
                            mime="application/pdf",
                            key="download_pdf",
                            use_container_width=True
                        )
                        
                except Exception as e:
                    st.error(f"❌ Error during PDF processing: {str(e)}")
                    st.markdown('<div class="error-message">Please try with a different PDF file or contact support if the issue persists.</div>', unsafe_allow_html=True)
                    
        except Exception as e:
            st.error(f"❌ Unexpected error: {str(e)}")
            st.markdown('<div class="error-message">The PDF file appears to be corrupted or incompatible. Please try with a different PDF file.</div>', unsafe_allow_html=True)

def display_pdf_preview(uploaded_file, pdf_reader):
    """Display PDF preview and information"""
    
    # PDF Preview Section
    st.markdown("---")
    st.subheader("📋 PDF Preview & Information")
    
    # Create columns for preview and info
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown('<div class="pdf-preview">', unsafe_allow_html=True)
        st.markdown("**📄 Document Preview**")
        
        # Get page count
        page_count = len(pdf_reader.pages)
        
        # Show page previews as numbered boxes
        st.write(f"**Total Pages:** {page_count}")
        
        # Display first 5 pages as preview
        preview_pages = min(5, page_count)
        if preview_pages > 0:
            cols = st.columns(preview_pages)
            
            for i in range(preview_pages):
                with cols[i]:
                    st.markdown(f'''
                    <div class="page-preview">
                        <div>
                            <h3>📄</h3>
                            <strong>Page {i+1}</strong>
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
        
        if page_count > 5:
            st.info(f"📖 Showing first 5 of {page_count} pages")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="file-info">', unsafe_allow_html=True)
        st.markdown("**📊 File Information**")
        
        # File size
        file_size = len(uploaded_file.getvalue())
        size_kb = file_size / 1024
        size_mb = size_kb / 1024
        
        if size_mb >= 1:
            st.write(f"**Size:** {size_mb:.1f} MB")
        else:
            st.write(f"**Size:** {size_kb:.1f} KB")
        
        # Page count
        st.write(f"**Pages:** {page_count}")
        
        # Get metadata
        try:
            metadata = pdf_reader.metadata
            if metadata:
                if '/Title' in metadata and metadata['/Title']:
                    st.write(f"**Title:** {metadata['/Title']}")
                if '/Author' in metadata and metadata['/Author']:
                    st.write(f"**Author:** {metadata['/Author']}")
                if '/Producer' in metadata and metadata['/Producer']:
                    st.write(f"**Producer:** {metadata['/Producer']}")
        except:
            pass
        
        st.write("**Status:** ✅ Ready to process")
        st.markdown('</div>', unsafe_allow_html=True)

def process_image():
    st.header("🖼️ Image Processing")
    
    uploaded_file = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif', 'webp'], key="image_uploader")
    
    if uploaded_file is not None:
        # Check if uploaded file is actually an image
        file_type = get_file_type(uploaded_file)
        if file_type != "image":
            st.markdown(f"""
            <div class="warning-message">
            <h4>⚠️ Wrong File Type Detected</h4>
            <p>You uploaded a <strong>{file_type.upper()}</strong> file, but this section is for <strong>IMAGE</strong> files.</p>
            <p><strong>Uploaded file:</strong> {uploaded_file.name}</p>
            <p>Please switch to the <strong>{file_type.capitalize()}</strong> section or upload an image file.</p>
            </div>
            """, unsafe_allow_html=True)
            return
            
        try:
            # Automatically display image preview
            image = Image.open(uploaded_file)
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.image(image, caption=f"Original Image - {uploaded_file.name}", use_column_width=True)
            
            with col2:
                st.markdown('<div class="file-info">', unsafe_allow_html=True)
                st.markdown("**📊 Image Information**")
                st.write(f"**Format:** {image.format}")
                st.write(f"**Size:** {image.size[0]} x {image.size[1]} pixels")
                st.write(f"**Mode:** {image.mode}")
                file_size = len(uploaded_file.getvalue()) / 1024
                st.write(f"**File Size:** {file_size:.1f} KB")
                st.write(f"**File Name:** {uploaded_file.name}")
                st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("🛠️ Image Processing Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**📐 Resize Options**")
                width = st.number_input("Width", min_value=1, value=image.width, key="image_width")
                height = st.number_input("Height", min_value=1, value=image.height, key="image_height")
                maintain_aspect = st.checkbox("Maintain Aspect Ratio", value=True, key="maintain_aspect")
                
                if maintain_aspect:
                    aspect_ratio = image.width / image.height
                    col1a, col1b = st.columns(2)
                    with col1a:
                        if st.button("Calculate Height from Width", key="calc_height"):
                            height = int(width / aspect_ratio)
                    with col1b:
                        if st.button("Calculate Width from Height", key="calc_width"):
                            width = int(height * aspect_ratio)
            
            with col2:
                st.markdown("**🎨 Format Options**")
                output_format = st.selectbox("Output Format", ["JPEG", "PNG", "BMP"], key="output_format")
                quality = st.slider("Quality (JPEG only)", 1, 100, 95, key="quality_slider")
                st.markdown("**💫 Enhancements**")
                enhance_contrast = st.checkbox("Enhance Contrast", key="enhance_contrast")
                convert_grayscale = st.checkbox("Convert to Grayscale", key="convert_grayscale")
            
            if st.button("🚀 Process Image", key="process_image_btn", type="primary"):
                try:
                    with st.spinner("🔄 Processing image..."):
                        # Apply transformations
                        processed_image = image.copy()
                        
                        # Resize
                        if maintain_aspect:
                            processed_image = processed_image.resize((width, height), Image.Resampling.LANCZOS)
                        else:
                            processed_image = processed_image.resize((width, height))
                        
                        # Enhancements
                        if enhance_contrast:
                            from PIL import ImageEnhance
                            enhancer = ImageEnhance.Contrast(processed_image)
                            processed_image = enhancer.enhance(1.5)
                        
                        if convert_grayscale:
                            processed_image = processed_image.convert('L')
                        
                        # Convert format if needed
                        if output_format == "JPEG":
                            if processed_image.mode in ('RGBA', 'LA', 'P'):
                                background = Image.new('RGB', processed_image.size, (255, 255, 255))
                                if processed_image.mode == 'P':
                                    processed_image = processed_image.convert('RGBA')
                                background.paste(processed_image, mask=processed_image.split()[-1] if processed_image.mode == 'RGBA' else None)
                                processed_image = background
                        
                        # Display processed image
                        st.image(processed_image, caption="Processed Image", use_column_width=True)
                        
                        # Prepare download
                        buffered = io.BytesIO()
                        if output_format == "JPEG":
                            processed_image.save(buffered, format="JPEG", quality=quality, optimize=True)
                        else:
                            processed_image.save(buffered, format=output_format, optimize=True)
                        
                        st.success("🎉 Image processed successfully!")
                        
                        # Download button
                        st.download_button(
                            label="📥 Download Processed Image",
                            data=buffered.getvalue(),
                            file_name=f"processed_image.{output_format.lower()}",
                            mime=f"image/{output_format.lower()}",
                            key="download_image",
                            use_container_width=True
                        )
                        
                except Exception as e:
                    st.error(f"❌ Error processing image: {str(e)}")
                    
        except Exception as e:
            st.error(f"❌ Error loading image: {str(e)}")

def main():
    # Initialize session state for ad clicks
    if 'ad_clicked' not in st.session_state:
        st.session_state.ad_clicked = None
    
    # Display all ads
    display_left_sidebar_ads()
    display_right_sidebar_ads()
    
    # Handle ad clicks
    handle_ad_clicks()
    
    # Main content with ads - get the middle column for main content
    middle_col = display_main_content_ads()
    
    # Main application content in the middle column
    with middle_col:
        st.title("📁 File Processing Application")
        st.write("Resize and process images, PDFs, CSV, DOCX, and Excel files")
        
        # File type selection
        file_type = st.selectbox(
            "Select File Type to Process",
            ["Image", "PDF", "CSV", "DOCX", "Excel"],
            key="file_type_main"
        )
        
        # Display appropriate processing function
        if file_type == "Image":
            process_image()
        elif file_type == "PDF":
            process_pdf()
        elif file_type == "CSV":
            process_csv()
        elif file_type == "DOCX":
            process_docx()
        elif file_type == "Excel":
            process_excel()

def process_csv():
    st.header("📊 CSV Processing")
    uploaded_file = st.file_uploader("Upload CSV", type=['csv'], key="csv_uploader")
    
    if uploaded_file is not None:
        # Check if uploaded file is actually a CSV
        file_type = get_file_type(uploaded_file)
        if file_type != "csv":
            st.markdown(f"""
            <div class="warning-message">
            <h4>⚠️ Wrong File Type Detected</h4>
            <p>You uploaded a <strong>{file_type.upper()}</strong> file, but this section is for <strong>CSV</strong> files.</p>
            <p><strong>Uploaded file:</strong> {uploaded_file.name}</p>
            <p>Please switch to the <strong>{file_type.capitalize()}</strong> section or upload a CSV file.</p>
            </div>
            """, unsafe_allow_html=True)
            return
            
        try:
            # Automatically display CSV preview
            df = pd.read_csv(uploaded_file)
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.write(f"**Dataset Shape:** {df.shape[0]} rows × {df.shape[1]} columns")
                st.dataframe(df.head(), use_container_width=True)
            
            with col2:
                st.markdown('<div class="file-info">', unsafe_allow_html=True)
                st.markdown("**📊 CSV Information**")
                st.write(f"**Rows:** {df.shape[0]:,}")
                st.write(f"**Columns:** {df.shape[1]}")
                file_size = len(uploaded_file.getvalue()) / 1024
                st.write(f"**File Size:** {file_size:.1f} KB")
                st.write(f"**File Name:** {uploaded_file.name}")
                st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("🛠️ CSV Processing Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**🧹 Data Cleaning**")
                remove_duplicates = st.checkbox("Remove duplicate rows", key="remove_duplicates")
                fill_na = st.checkbox("Fill missing values", key="fill_na")
                if fill_na:
                    fill_method = st.selectbox("Fill method", ["Forward Fill", "Backward Fill", "Mean", "Median", "Custom"], key="fill_method")
                    if fill_method == "Custom":
                        custom_value = st.text_input("Custom value", key="custom_value")
            
            with col2:
                st.markdown("**📊 Column Operations**")
                columns_to_keep = st.multiselect("Select columns to keep", df.columns.tolist(), default=df.columns.tolist(), key="columns_to_keep")
                rename_cols = st.checkbox("Rename columns", key="rename_cols")
                if rename_cols:
                    new_names = st.text_area("New column names (comma-separated)", key="new_names")
            
            if st.button("🚀 Process CSV", key="process_csv_btn", type="primary"):
                try:
                    with st.spinner("🔄 Processing CSV..."):
                        processed_df = df.copy()
                        
                        # Keep selected columns
                        processed_df = processed_df[columns_to_keep]
                        
                        # Remove duplicates
                        if remove_duplicates:
                            initial_rows = len(processed_df)
                            processed_df = processed_df.drop_duplicates()
                            removed_count = initial_rows - len(processed_df)
                            if removed_count > 0:
                                st.info(f"✅ Removed {removed_count} duplicate rows")
                        
                        # Fill missing values
                        if fill_na:
                            if fill_method == "Forward Fill":
                                processed_df = processed_df.ffill()
                            elif fill_method == "Backward Fill":
                                processed_df = processed_df.bfill()
                            elif fill_method == "Mean":
                                processed_df = processed_df.fillna(processed_df.mean(numeric_only=True))
                            elif fill_method == "Median":
                                processed_df = processed_df.fillna(processed_df.median(numeric_only=True))
                            elif fill_method == "Custom" and custom_value:
                                processed_df = processed_df.fillna(custom_value)
                        
                        # Rename columns
                        if rename_cols and new_names:
                            new_cols = [name.strip() for name in new_names.split(',')]
                            if len(new_cols) == len(processed_df.columns):
                                processed_df.columns = new_cols
                                st.info(f"✅ Renamed {len(new_cols)} columns")
                        
                        st.success("🎉 CSV processed successfully!")
                        
                        # Show processed data
                        st.write("**Processed Data Preview:**")
                        st.dataframe(processed_df.head(), use_container_width=True)
                        st.write(f"**Final Shape:** {processed_df.shape[0]} rows × {processed_df.shape[1]} columns")
                        
                        # Download
                        csv_buffer = io.BytesIO()
                        processed_df.to_csv(csv_buffer, index=False)
                        csv_buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Download Processed CSV",
                            data=csv_buffer.getvalue(),
                            file_name="processed_data.csv",
                            mime="text/csv",
                            key="download_csv",
                            use_container_width=True
                        )
                        
                except Exception as e:
                    st.error(f"❌ Error processing CSV: {str(e)}")
                    
        except Exception as e:
            st.error(f"❌ Error loading CSV: {str(e)}")

def process_docx():
    st.header("📝 DOCX Processing")
    uploaded_file = st.file_uploader("Upload DOCX", type=['docx'], key="docx_uploader")
    
    if uploaded_file is not None:
        # Check if uploaded file is actually a DOCX
        file_type = get_file_type(uploaded_file)
        if file_type != "docx":
            st.markdown(f"""
            <div class="warning-message">
            <h4>⚠️ Wrong File Type Detected</h4>
            <p>You uploaded a <strong>{file_type.upper()}</strong> file, but this section is for <strong>DOCX</strong> files.</p>
            <p><strong>Uploaded file:</strong> {uploaded_file.name}</p>
            <p>Please switch to the <strong>{file_type.capitalize()}</strong> section or upload a DOCX file.</p>
            </div>
            """, unsafe_allow_html=True)
            return
            
        try:
            doc = Document(uploaded_file)
            
            # Display document info
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.success("✅ DOCX loaded successfully!")
                st.subheader("Document Content Preview")
                
                # Extract and display text content
                full_text = []
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        full_text.append(para.text)
                    if i >= 20:  # Limit preview to first 20 paragraphs
                        break
                
                preview_text = "\n".join(full_text)
                st.text_area("Content", preview_text, height=200, key="doc_content_preview")
                
                if len(doc.paragraphs) > 20:
                    st.info(f"📄 Showing first 20 of {len(doc.paragraphs)} paragraphs")
            
            with col2:
                st.markdown('<div class="file-info">', unsafe_allow_html=True)
                st.markdown("**📊 Document Information**")
                st.write(f"**File Name:** {uploaded_file.name}")
                file_size = len(uploaded_file.getvalue()) / 1024
                st.write(f"**File Size:** {file_size:.1f} KB")
                st.write(f"**Paragraphs:** {len(doc.paragraphs)}")
                st.write(f"**Sections:** {len(doc.sections)}")
                st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("🛠️ DOCX Processing Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**🔍 Text Operations**")
                find_text = st.text_input("Find text", key="find_text")
                replace_text = st.text_input("Replace with", key="replace_text")
            
            with col2:
                st.markdown("**📤 Export Options**")
                export_format = st.selectbox("Export as", ["DOCX", "TXT"], key="export_format")
            
            if st.button("🚀 Process Document", key="process_doc_btn", type="primary"):
                try:
                    with st.spinner("🔄 Processing document..."):
                        new_doc = Document()
                        
                        processed_paragraphs = 0
                        for para in doc.paragraphs:
                            text = para.text
                            if find_text and replace_text:
                                text = text.replace(find_text, replace_text)
                            
                            if text.strip():
                                new_doc.add_paragraph(text)
                                processed_paragraphs += 1
                        
                        st.success(f"🎉 Document processed successfully! Processed {processed_paragraphs} paragraphs.")
                        
                        if export_format == "DOCX":
                            buffer = io.BytesIO()
                            new_doc.save(buffer)
                            buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Download Processed DOCX",
                                data=buffer.getvalue(),
                                file_name="processed_document.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_docx",
                                use_container_width=True
                            )
                        else:
                            full_text = []
                            for para in new_doc.paragraphs:
                                full_text.append(para.text)
                            
                            text_content = "\n".join(full_text)
                            st.download_button(
                                label="📥 Download as TXT",
                                data=text_content,
                                file_name="processed_document.txt",
                                mime="text/plain",
                                key="download_txt",
                                use_container_width=True
                            )
                            
                except Exception as e:
                    st.error(f"❌ Error processing document: {str(e)}")
                    
        except Exception as e:
            st.error(f"❌ Error loading DOCX: {str(e)}")

def process_excel():
    st.header("📈 Excel Processing")
    uploaded_file = st.file_uploader("Upload Excel", type=['xlsx', 'xls'], key="excel_uploader")
    
    if uploaded_file is not None:
        # Check if uploaded file is actually an Excel file
        file_type = get_file_type(uploaded_file)
        if file_type != "excel":
            st.markdown(f"""
            <div class="warning-message">
            <h4>⚠️ Wrong File Type Detected</h4>
            <p>You uploaded a <strong>{file_type.upper()}</strong> file, but this section is for <strong>EXCEL</strong> files.</p>
            <p><strong>Uploaded file:</strong> {uploaded_file.name}</p>
            <p>Please switch to the <strong>{file_type.capitalize()}</strong> section or upload an Excel file.</p>
            </div>
            """, unsafe_allow_html=True)
            return
            
        try:
            xl_file = pd.ExcelFile(uploaded_file)
            sheet_names = xl_file.sheet_names
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.success(f"✅ Excel file loaded successfully!")
                st.write(f"**Sheets found:** {', '.join(sheet_names)}")
                
                selected_sheet = st.selectbox("Select Sheet to Preview", sheet_names, key="selected_sheet")
                
                if selected_sheet:
                    df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
                    st.write(f"**Sheet:** {selected_sheet} - {df.shape[0]} rows × {df.shape[1]} columns")
                    st.dataframe(df.head(), use_container_width=True)
            
            with col2:
                st.markdown('<div class="file-info">', unsafe_allow_html=True)
                st.markdown("**📊 Excel Information**")
                st.write(f"**File Name:** {uploaded_file.name}")
                file_size = len(uploaded_file.getvalue()) / 1024
                st.write(f"**File Size:** {file_size:.1f} KB")
                st.write(f"**Sheets:** {len(sheet_names)}")
                st.write(f"**Sheets:** {', '.join(sheet_names)}")
                st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("🛠️ Excel Processing Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**📊 Data Operations**")
                operations = st.multiselect("Select operations", 
                                          ["Remove Duplicates", "Sort Data", "Filter Data"], 
                                          key="operations")
                
                if "Sort Data" in operations:
                    if selected_sheet:
                        df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
                        sort_column = st.selectbox("Sort by column", df.columns.tolist(), key="sort_column")
                        sort_ascending = st.checkbox("Ascending", value=True, key="sort_ascending")
                
                if "Filter Data" in operations:
                    if selected_sheet:
                        df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
                        filter_column = st.selectbox("Filter by column", df.columns.tolist(), key="filter_column")
                        filter_value = st.text_input("Filter value", key="filter_value")
            
            with col2:
                st.markdown("**📤 Export Options**")
                output_sheets = st.multiselect("Sheets to include", sheet_names, default=sheet_names, key="output_sheets")
                output_format = st.selectbox("Output Format", ["Excel", "CSV"], key="excel_output_format")
            
            if st.button("🚀 Process Excel", key="process_excel_btn", type="primary"):
                try:
                    with st.spinner("🔄 Processing Excel file..."):
                        processed_data = {}
                        
                        for sheet in output_sheets:
                            sheet_df = pd.read_excel(uploaded_file, sheet_name=sheet)
                            
                            # Apply operations
                            if "Remove Duplicates" in operations:
                                initial_rows = len(sheet_df)
                                sheet_df = sheet_df.drop_duplicates()
                                removed_count = initial_rows - len(sheet_df)
                                if removed_count > 0:
                                    st.info(f"✅ Removed {removed_count} duplicate rows from {sheet}")
                            
                            if "Sort Data" in operations and sort_column in sheet_df.columns:
                                sheet_df = sheet_df.sort_values(by=sort_column, ascending=sort_ascending)
                                st.info(f"✅ Sorted {sheet} by {sort_column}")
                            
                            if "Filter Data" in operations and filter_column in sheet_df.columns:
                                initial_rows = len(sheet_df)
                                sheet_df = sheet_df[sheet_df[filter_column].astype(str).str.contains(filter_value, na=False)]
                                filtered_count = initial_rows - len(sheet_df)
                                if filtered_count > 0:
                                    st.info(f"✅ Filtered {filtered_count} rows from {sheet}")
                            
                            processed_data[sheet] = sheet_df
                        
                        st.success("🎉 Excel file processed successfully!")
                        
                        if output_format == "Excel":
                            buffer = io.BytesIO()
                            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                                for sheet_name, sheet_df in processed_data.items():
                                    sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)
                            
                            buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Download Processed Excel",
                                data=buffer.getvalue(),
                                file_name="processed_data.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key="download_excel",
                                use_container_width=True
                            )
                        else:
                            # Export first sheet as CSV
                            first_sheet = list(processed_data.keys())[0]
                            csv_buffer = io.BytesIO()
                            processed_data[first_sheet].to_csv(csv_buffer, index=False)
                            csv_buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Download as CSV",
                                data=csv_buffer.getvalue(),
                                file_name="processed_data.csv",
                                mime="text/csv",
                                key="download_excel_csv",
                                use_container_width=True
                            )
                            
                except Exception as e:
                    st.error(f"❌ Error processing Excel: {str(e)}")
                    
        except Exception as e:
            st.error(f"❌ Error loading Excel: {str(e)}")

def parse_page_range(page_range_str, max_pages):
    """Parse page range string like '1-3,5,7-9'"""
    pages = set()
    ranges = page_range_str.split(',')
    
    for r in ranges:
        r = r.strip()
        if '-' in r:
            start, end = r.split('-')
            try:
                start = int(start.strip())
                end = int(end.strip())
                # Ensure valid range
                if start > end:
                    start, end = end, start
                pages.update(range(max(1, start), min(max_pages, end) + 1))
            except ValueError:
                st.error(f"Invalid range: {r}")
        else:
            try:
                page_num = int(r)
                if 1 <= page_num <= max_pages:
                    pages.add(page_num)
            except ValueError:
                st.error(f"Invalid page number: {r}")
    
    return sorted(list(pages))

if __name__ == "__main__":
    main()